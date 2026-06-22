from docx import Document
from docx.document import Document as _Document

from logic.docx_tools import (
    format_kazakh_typography,
    set_run_typography,
    set_paragraph_tightness,
)


class GrowCardBuilder:
    def __init__(self, source):
        if isinstance(source, _Document):
            self.docx = source
        else:
            # If str or BytesIO is given, python-docx will handle it itself
            self.docx = Document(source)

    def build(
        self, students_cards: list[dict], academic_year: str, group_name: str
    ) -> _Document:
        self._keep_only_first_page_header(self.docx)

        for index, student in enumerate(students_cards):
            if index > 0:
                self.docx.add_page_break()

            self._add_title(
                self.docx,
                f"{academic_year} оқу жылына арналған баланың жеке даму картасы",
            )
            fullname = format_kazakh_typography(student["fullname"])
            birth_date = student["birth_date"].strip()
            birth_date = f"{birth_date} ж." if birth_date else ""

            self._add_meta_paragraph(self.docx, "Баланың Т.А.Ә.:", fullname)
            self._add_meta_paragraph(self.docx, "Баланың туған жылы, күні:", birth_date)
            self._add_meta_paragraph(self.docx, "Топ:", group_name)

            self._create_assessment_table(self.docx, student["assessments"])
        return self.docx

    def _keep_only_first_page_header(self, docx):
        target_p_idx = -1
        for idx, p in enumerate(docx.paragraphs):
            if "баланың т.а.ә" in p.text.lower() or "т.а.ә" in p.text.lower():
                target_p_idx = idx
                break

        if target_p_idx != -1:
            start_p_remove = max(0, target_p_idx - 2)
            for p_idx in range(len(docx.paragraphs) - 1, start_p_remove - 1, -1):
                p_element = docx.paragraphs[p_idx]._p
                p_element.getparent().remove(p_element)

        while len(docx.tables) > 0:
            tbl_element = docx.tables[0]._element
            tbl_element.getparent().remove(tbl_element)

    def _add_title(self, docx, title):
        p = docx.add_paragraph()
        p.alignment = 1
        set_paragraph_tightness(p)
        run = p.add_run(title)
        run.bold = True
        set_run_typography(run, size_pt=12)

    def _add_meta_paragraph(self, docx, label: str, value: str):
        p = docx.add_paragraph()
        set_paragraph_tightness(p)

        run_label = p.add_run(f"{label} ")
        run_label.bold = True
        set_run_typography(run_label, size_pt=12)

        run_value = p.add_run(value)
        run_value.bold = False
        set_run_typography(run_value, size_pt=12)

    def _create_assessment_table(self, docx, assessments: list[dict]):
        rows_count = len(assessments) + 1
        table = docx.add_table(rows=rows_count, cols=5)
        table.style = "Table Grid"
        headers = [
            "Құзыреттіліктер",
            "Бастапқы бақылау нәтижелері бойынша дамыту, түзету іс-шаралары (қазан-желтоқсан)",
            "Аралық бақылау нәтижелері бойынша дамыту, түзету ісшаралары (ақпан-сәуір)",
            "Қорытынды бақылау нәтижелері бойынша дамыту, түзету ісшаралары (маусым-тамыз)",
            (
                "Қорытынды (баланың даму деңгейі сәйкес келеді:"
                " III деңгей - «жоғары»; II деңгей – «орташа»; I деңгей - «төмен»)"
            ),
        ]
        hdr_cells = table.rows[0].cells
        for i, header_text in enumerate(headers):
            hdr_cells[i].text = header_text
            for p in hdr_cells[i].paragraphs:
                p.alignment = 1
                for run in p.runs:
                    run.bold = True
                    set_run_typography(run, size_pt=12)
        for row_idx, assess in enumerate(assessments, start=1):
            row_cells = table.rows[row_idx].cells

            row_cells[0].text = assess["criterion"]
            row_cells[1].text = format_kazakh_typography(assess["start"])
            row_cells[2].text = format_kazakh_typography(assess["mid"])
            row_cells[3].text = format_kazakh_typography(assess["end"])
            row_cells[4].text = ""

            for cell in row_cells:
                for p in cell.paragraphs:
                    set_paragraph_tightness(p)
                    for run in p.runs:
                        set_run_typography(run, size_pt=12)
