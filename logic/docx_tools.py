import re

from docx import Document
from docx.document import Document as _Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import Pt


def set_run_typography(run, font_name="Times New Roman", size_pt=14):
    run.font.name = font_name
    run.font.size = Pt(size_pt)


def format_kazakh_typography(text: str) -> str:
    if not text:
        return ""
    # Remove extra spaces before commas, periods, and colons
    text = re.sub(r"\s+([,.:;])", r"\1", text)
    # Add a space after the symbols (if none exist)
    text = re.sub(r"([,.:;])(?![\s$])", r"\1 ", text)
    # Replace straight quotes with Kazakh "«»"
    text = re.sub(r'"([^"]+)"', r"«\1»", text)
    # Clean up excess spaces
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def set_paragraph_tightness(paragraph, space_before=0, space_after=0, line_spacing=1.0):
    p_format = paragraph.paragraph_format
    p_format.space_before = Pt(space_before)
    p_format.space_after = Pt(space_after)
    p_format.line_spacing = line_spacing


def analyze_template_placeholders(template_path, data_keys):
    doc = Document(template_path)
    template_keys = set()
    pattern = re.compile(r"\{\{([^}]+)\}\}")

    for paragraph in doc.paragraphs:
        for match in pattern.finditer(paragraph.text):
            template_keys.add(match.group(1).strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for match in pattern.finditer(paragraph.text):
                        template_keys.add(match.group(1).strip())

    expected_set = set(data_keys)

    # Keys NOT in the template (but present in the data)
    missing_in_template = expected_set - template_keys

    # EXTRA keys in the template (but not in the data)
    extra_in_template = template_keys - expected_set

    # Intersection of two sides (Exactly exchangeable keys)
    intersection_keys = template_keys & expected_set

    return missing_in_template, extra_in_template, intersection_keys


def replace_placeholders_in_document(element, replacements):
    if hasattr(element, "paragraphs"):
        for paragraph in element.paragraphs:
            for key, value in replacements.items():
                placeholder = f"{{{{{key}}}}}"
                for run in paragraph.runs:
                    run.text = run.text.replace(placeholder, value)
    if hasattr(element, "tables"):
        for table in element.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_placeholders_in_document(cell, replacements)


def create_children_grow_cards(template_path, children_data, progress_callback=None):
    if not children_data:
        raise ValueError("Қате: Өңдеуге жіберілген балалар деректер жиыны бос!")
    all_data_keys = {key for child in children_data for key in child.keys()}
    missing, extra, _ = analyze_template_placeholders(template_path, all_data_keys)
    if extra:
        raise ValueError(
            f"Қате: Шаблонда белгісіз маркерлер бар: {', '.join(extra)}. "
            f"Оларды Word файлынан өшіріңіз немесе баптаудан шаблонға сәйкес қылыңыз!"
        )
    if missing:
        raise ValueError(
            f"Қате: Мына деректер шаблонда қолданылмаған: {', '.join(missing)}. "
            f"Шаблонға тиісті маркерлерді қосыңыз немесе баптаудан шаблонға сәйкес қылыңыз!"
        )

    merged_doc = Document(template_path)
    replace_placeholders_in_document(merged_doc, children_data[0])
    merged_body = merged_doc.element.body
    total_children = len(children_data)

    if progress_callback:
        progress_callback(children_data[0]["fullname"], 1, total_children)

    for index, child_data in enumerate(children_data[1:]):
        if progress_callback:
            progress_callback(child_data["fullname"], index + 2, total_children)
        template_doc = Document(template_path)
        replace_placeholders_in_document(template_doc, child_data)
        merged_doc.add_page_break()

        sect_pr = merged_body.xpath("w:sectPr")[0]
        for element in template_doc.element.body:
            if "sectPr" in element.tag:
                continue
            if isinstance(element, CT_P) and not element.text.strip():
                continue
            merged_body.insert(merged_body.index(sect_pr), element)
    return merged_doc


def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def fill_specific_child_in_big_file(doc, target_child_name, values, col_index=2):
    is_fully_processed = False
    found_child_paragraph = False

    for block in iter_block_items(doc):
        # Search for a child's name (Paragraph)
        if isinstance(block, Paragraph) and "Баланың Т.А.Ә" in block.text:
            if target_child_name in block.text:
                found_child_paragraph = True
            continue

        if isinstance(block, Table) and found_child_paragraph:
            if not block.rows or "Құзыреттіліктер" not in block.cell(0, 0).text:
                continue

            # Table populating logic
            for i, val in enumerate(values, start=1):
                try:
                    block.cell(i, col_index).text = str(val)
                except IndexError:
                    pass

            is_fully_processed = True
            found_child_paragraph = False
            break

    return is_fully_processed


def fill_all_children_in_big_file(
    docx_path, children_data, col_index, progress_callback=None
):
    doc = Document(docx_path)
    missing_children = []

    for index, child in enumerate(children_data, start=1):
        child_name = child["fullname"]
        values = [v for k, v in child.items() if k != "fullname"]
        if progress_callback:
            progress_callback(child_name, index, len(children_data))

        success = fill_specific_child_in_big_file(doc, child_name, values, col_index)

        if not success:
            missing_children.append(child_name)

    return doc, missing_children
